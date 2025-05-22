import socket
import ssl
import requests
from urllib.parse import urlparse
from datetime import datetime
from docx import Document
import subprocess

# Input and output files
WEBSITE_LIST_FILE = "websites.txt"
OUTPUT_FILE = "website_check_log.docx"

def read_websites(file_path):
    with open(file_path, 'r') as file:
        # Add https:// if not present
        return [("https://" + line.strip() if not line.strip().startswith(('http://', 'https://')) else line.strip()) for line in file if line.strip()]

def nslookup(domain):
    try:
        result = subprocess.check_output(['nslookup', domain], universal_newlines=True)
        return result
    except subprocess.CalledProcessError:
        return None

def check_dns(domain):
    try:
        ip = socket.gethostbyname(domain)
        return ip
    except socket.gaierror:
        return None

def check_http(site):
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/122.0.0.0 Safari/537.36"
        )
    }
    try:
        response = requests.get(site, headers=headers, timeout=5)
        if response.status_code == 200:
            return "HTTP OK"
        else:
            # Instead of error, mark as working but HTTP issue
            return f"Working, but with HTTP Error: Status Code {response.status_code}"
    except requests.RequestException as e:
        return f"HTTP Error: {e}"

def check_ssl(domain):
    context = ssl.create_default_context()
    try:
        with socket.create_connection((domain, 443), timeout=5) as sock:
            with context.wrap_socket(sock, server_hostname=domain) as ssock:
                cert = ssock.getpeercert()
                # Extract the certificate expiry date
                cert_expiry = cert['notAfter']
                cert_expiry = datetime.strptime(cert_expiry, "%b %d %H:%M:%S %Y GMT")
                
                # Calculate the remaining days until certificate expiry
                remaining_days = (cert_expiry - datetime.utcnow()).days
                return f"SSL Valid, Expires on {cert_expiry.strftime('%Y-%m-%d %H:%M:%S')}, {remaining_days} days remaining"
    except Exception as e:
        return f"SSL Error: {e}"

def write_to_docx(results):
    doc = Document()
    doc.add_heading('Website Health Check Report', 0)

    for i, result in enumerate(results, start=1):
        doc.add_paragraph(f"{i}. {result['domain']} --> {result['status']}")
        doc.add_paragraph(f"   Reason: {result['reason']}")
        doc.add_paragraph(f"   DNS Lookup:\n{result['dns_lookup']}")
        doc.add_paragraph(f"   SSL Certificate: {result['ssl_cert']}")

    doc.save(OUTPUT_FILE)
    print(f"\nResults written to {OUTPUT_FILE}")

def main():
    websites = read_websites(WEBSITE_LIST_FILE)
    results = []

    for site in websites:
        parsed = urlparse(site)
        domain = parsed.netloc or parsed.path

        result = {
            "domain": site,
            "status": "",
            "reason": "",
            "dns_lookup": "",
            "ssl_cert": "",
        }

        # Check DNS
        dns_ip = check_dns(domain)
        if not dns_ip:
            result["status"] = "DNS Failed"
            result["reason"] = "Could not resolve DNS."
            result["dns_lookup"] = nslookup(domain) if nslookup(domain) else "No nslookup info available"
        else:
            result["status"] = "working"
            result["reason"] = "No issues found."
            result["dns_lookup"] = nslookup(domain) if nslookup(domain) else "No nslookup info available"

            # Check HTTP status
            http_status = check_http(site)
            if http_status != "HTTP OK":
                result["status"] = "working"
                result["reason"] = f"Working, but with HTTP error: {http_status}"
            elif site.startswith("https://"):
                # Check SSL details
                ssl_status = check_ssl(domain)
                result["ssl_cert"] = ssl_status
                if "SSL Error" in ssl_status:
                    result["status"] = "SSL Failed"
                    result["reason"] = ssl_status
                else:
                    result["status"] = "working"
                    result["reason"] = "No issues found."

        results.append(result)

    write_to_docx(results)

if __name__ == "__main__":
    main()
